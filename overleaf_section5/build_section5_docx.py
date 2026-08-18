#!/usr/bin/env python3
"""Build a visually reviewed Word copy of the Section 5 LaTeX source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "section5_zero_shot_predictor.tex"
OUTPUT = ROOT / "Section_5_Risk_Calibrated_Predictor_Review_v3.docx"

INK = RGBColor(36, 45, 58)
BLUE = RGBColor(31, 77, 120)
ORANGE = RGBColor(196, 103, 0)
MUTED = RGBColor(91, 100, 114)
LIGHT_BLUE = "E8EEF5"
LIGHT_ORANGE = "FFF4E5"
HEADER_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF6EE"
PALE_GRAY = "F5F7FA"
POSITIVE = RGBColor(24, 112, 68)
NEGATIVE = RGBColor(166, 49, 49)
WHITE = "FFFFFF"

EQ_NUMBERS = {
    "eq:risk-target": 1,
    "eq:harm-event": 2,
    "eq:risk-score": 3,
    "eq:risk-selection": 4,
}

REFS = {
    "sec:real-context": "3",
    "sec:context-mechanisms": "4",
    "sec:risk-target": "5.1",
    "sec:risk-calibration": "5.2",
    "sec:risk-transfer": "5.3",
    "sec:risk-datasets": "5.4",
    "sec:risk-limits": "5.5",
    "sec:risk-implications": "5.6",
    "tab:risk-profiles": "1",
    "tab:risk-models": "2",
    "tab:risk-datasets": "3",
    "fig:risk-pipeline": "1",
    "fig:compute-harm-dial": "2",
    "fig:risk-models": "3",
    "eq:normalized-mase": "the normalized-MASE definition in Section 3",
    **{label: f"({number})" for label, number in EQ_NUMBERS.items()},
}

EQUATIONS = {
    "eq:risk-target": "rᵢ,ₖ,ₕ = log(Eᵢ,ₖ,ₕ / Eᵢ,native,ₕ)",
    "eq:harm-event": "bᵢ,ₖ,ₕ = 𝟙[Eᵢ,ₖ,ₕ > 1.05 · Eᵢ,native,ₕ]",
    "eq:risk-score": "qᵢ,ₖ,ₕ = μ̂ᵢ,ₖ,ₕ + u · σ̂ᵢ,ₖ,ₕ + v · p̂⁽⁵⁾ᵢ,ₖ,ₕ",
    "eq:risk-selection": (
        "L̂⁽ᵃ⁾ᵢ,ₕ = min{Lₖ : qᵢ,ₖ,ₕ ≤ τₐ},  or Lnative if the set is empty"
    ),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{side}"))
        if tag is None:
            tag = OxmlElement(f"w:{side}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table) -> None:
    """Use restrained academic rules instead of a dense spreadsheet grid."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    specs = {
        "top": ("single", "8", HEADER_BLUE),
        "bottom": ("single", "8", HEADER_BLUE),
        "start": ("nil", "0", WHITE),
        "end": ("nil", "0", WHITE),
        "insideH": ("single", "3", "C9D2DC"),
        "insideV": ("nil", "0", WHITE),
    }
    for edge, (value, size, color) in specs.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        element.set(qn("w:space"), "0")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    normal.font.size = Pt(10.3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.16

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption.font.bold = False
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.05

    eq_style = styles.add_style("Display Equation", WD_STYLE_TYPE.PARAGRAPH)
    eq_style.font.name = "Cambria Math"
    eq_style._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    eq_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    eq_style.font.size = Pt(10.5)
    eq_style.font.color.rgb = INK
    eq_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_style.paragraph_format.space_before = Pt(7)
    eq_style.paragraph_format.space_after = Pt(9)
    eq_style.paragraph_format.keep_together = True

    note = styles.add_style("Review Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    note.font.name = "Calibri"
    note.font.size = Pt(10.5)
    note.font.color.rgb = MUTED
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(16)


def configure_page(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PREDICTCSL  ·  SECTION 5 REVIEW COPY")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Page ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    add_field(p, "PAGE")


def strip_comments(tex: str) -> str:
    return re.sub(r"(?<!\\)%.*$", "", tex, flags=re.MULTILINE)


def unwrap_commands(text: str) -> str:
    patterns = [
        r"\\(?:emph|textbf|mathbf|boldsymbol|text|operatorname)\{([^{}]*)\}",
        r"\\widehat\{([^{}]*)\}",
        r"\\min\{([^{}]*)\}",
    ]
    previous = None
    while previous != text:
        previous = text
        for pattern in patterns:
            text = re.sub(pattern, r"\1", text)
    return text


def plain(text: str) -> str:
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\ref\{([^}]+)\}", lambda m: REFS.get(m.group(1), "?"), text)
    math_replacements = {
        r"h\in\{16,32,64,128,512,1024\}": "h ∈ {16, 32, 64, 128, 512, 1024}",
        r"\widehat\mu_{i,k,h}": "μ̂ᵢ,ₖ,ₕ",
        r"\widehat\sigma_{i,k,h}": "σ̂ᵢ,ₖ,ₕ",
        r"\widehat p^{(5)}_{i,k,h}": "p̂⁽⁵⁾ᵢ,ₖ,ₕ",
        r"E_{i,\mathrm{native},h}": "Eᵢ,native,ₕ",
        r"E_{i,k,h}": "Eᵢ,ₖ,ₕ",
        r"L_k": "Lₖ",
        r"\tau_a": "τₐ",
        r"\widehat L_{i,h}^{(a)}": "L̂⁽ᵃ⁾ᵢ,ₕ",
        r"\widehat{\boldsymbol{s}}_{i,h}": "ŝᵢ,ₕ",
        r"\widehat{L}_c": "L̂c",
        r"\widehat L_c": "L̂c",
        r"\widehat{k}_c": "k̂c",
        r"\widehat k_c": "k̂c",
        r"L_k\in\mathcal{G}_m": "Lₖ ∈ Gₘ",
        r"\mathcal{G}_m": "Gₘ",
        r"\mathcal{V}_{i,h}": "Vᵢ,ₕ",
        r"\mathcal{V}_c": "Vc",
    }
    for old, new in math_replacements.items():
        text = text.replace(old, new)
    text = unwrap_commands(text)
    replacements = {
        r"\%": "%", r"\&": "&", r"\_": "_", r"\{": "{", r"\}": "}", "~": " ",
        "--": "–", r"\Delta": "Δ", r"\lambda": "λ", r"\mu": "μ",
        r"\sigma": "σ", r"\widehat": "", r"\arg": "arg ",
        r"\min": "min", r"\in": "∈", r"\sum": "∑", r"\times": "×",
        r"\mathrm": "", r"\mathcal": "", r"\boldsymbol": "",
        r"\left": "", r"\right": "", r"\qquad": "    ",
        r"\,": " ", r"\!": "", r"\\": " ", "$": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace("Equation the normalized-MASE definition in Section 3",
                        "the normalized-MASE definition in Section 3")
    text = text.replace(
        "Accuracy is the normalized MASE in the normalized-MASE definition in Section 3.",
        "Accuracy follows the normalized-MASE definition in Section 3.",
    )
    return text


def extract_braced(text: str, command: str) -> str:
    marker = f"\\{command}{{"
    start = text.find(marker)
    if start < 0:
        return ""
    pos = start + len(marker)
    depth = 1
    for idx in range(pos, len(text)):
        if text[idx] == "{" and text[idx - 1] != "\\":
            depth += 1
        elif text[idx] == "}" and text[idx - 1] != "\\":
            depth -= 1
            if depth == 0:
                return text[pos:idx]
    raise ValueError(f"Unclosed {command} block")


def add_body_paragraph(doc: Document, text: str) -> None:
    cleaned = plain(text)
    if not cleaned:
        return
    p = doc.add_paragraph(style="Normal")
    p.add_run(cleaned)


def add_equation(doc: Document, block: str) -> None:
    label_match = re.search(r"\\label\{([^}]+)\}", block)
    if not label_match:
        raise ValueError("Equation without label")
    label = label_match.group(1)
    p = doc.add_paragraph(style="Display Equation")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.65))
    run = p.add_run(EQUATIONS[label])
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    p.add_run(f"    ({EQ_NUMBERS[label]})")


def add_figure(doc: Document, block: str) -> None:
    path_match = re.search(
        r"\\includegraphics(?:\[[^]]*\])?\s*\{([^}]+)\}", block
    )
    label_match = re.search(r"\\label\{([^}]+)\}", block)
    if not path_match or not label_match:
        raise ValueError("Incomplete figure block")
    image_path = (ROOT / path_match.group(1)).with_suffix(".png")
    label = label_match.group(1)
    width = {
        "fig:risk-pipeline": 6.75,
        "fig:compute-harm-dial": 6.15,
        "fig:risk-models": 5.55,
    }[label]
    alt_text = {
        "fig:risk-pipeline": (
            "Two-row workflow: synthetic TSFM labels train an expected-risk "
            "model and calibrate five harm profiles; at deployment, a chosen "
            "profile selects the shortest acceptable window or native fallback."
        ),
        "fig:compute-harm-dial": (
            "Two-panel chart across eleven forecasters: five ordered profiles "
            "trade pooled incidence of more than five percent harm for mean "
            "theoretical FLOPs saved while aggregate MASE remains near native."
        ),
        "fig:risk-models": (
            "Two heatmaps across eleven forecasters and five profiles showing "
            "theoretical FLOPs saved and observed incidence of more than five "
            "percent harm; the FlowState row is outlined as a calibration miss."
        ),
    }[label]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", f"Figure {REFS[label]}")
    caption = plain(extract_braced(block, "caption"))
    cap = doc.add_paragraph(style="Caption")
    prefix = cap.add_run(f"Figure {REFS[label]}. ")
    prefix.bold = True
    cap.add_run(caption)


def parse_table_rows(block: str) -> list[list[str]]:
    tabular = re.search(r"\\begin\{tabular\}\{[^}]+\}(.*?)\\end\{tabular\}",
                        block, flags=re.DOTALL)
    if not tabular:
        raise ValueError("Missing tabular data")
    content = tabular.group(1)
    content = re.sub(r"\\(?:toprule|midrule|bottomrule)", "", content)
    content = re.sub(r"\\cmidrule(?:\([^)]*\))?\{[^}]+\}", "", content)
    rows = []
    for raw in re.split(r"\\\\", content):
        raw = raw.strip()
        if not raw:
            continue
        rows.append([plain(cell) for cell in raw.split("&")])
    return rows


def add_table(doc: Document, block: str) -> None:
    label_match = re.search(r"\\label\{([^}]+)\}", block)
    if not label_match:
        raise ValueError("Table without label")
    label = label_match.group(1)
    caption = plain(extract_braced(block, "caption"))
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True
    prefix = cap.add_run(f"Table {REFS[label]}. ")
    prefix.bold = True
    cap.add_run(caption)

    rows = parse_table_rows(block)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    if label == "tab:risk-profiles":
        widths = [1880, 1180, 1330, 1080, 1300, 1370, 1140]
    elif label == "tab:risk-datasets":
        widths = [2860, 1540, 1640, 1460, 1780]
    else:
        widths = [2300, 1350, 1350, 1500, 1400, 1380]
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for row_idx, (row, values) in enumerate(zip(table.rows, rows)):
        for col_idx, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(8.2 if row_idx else 8.4)
            run.font.color.rgb = INK
            if row_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, HEADER_BLUE)
            else:
                if label == "tab:risk-profiles":
                    fills = {
                        "Conservative": "E8F5F1",
                        "Balanced": "EAF2FF",
                        "Aggressive": "FFF3E0",
                        "Efficiency": "FDECEA",
                        "Max efficiency": "F3E8F2",
                    }
                    fill = fills.get(values[0], WHITE if row_idx % 2 else PALE_GRAY)
                    set_cell_shading(cell, fill)
                    if col_idx == 0:
                        run.bold = True
                    if col_idx == 6:
                        run.font.color.rgb = (POSITIVE if values[col_idx].startswith("+")
                                              else NEGATIVE)
                elif label == "tab:risk-models":
                    improved = values[5].startswith("+")
                    set_cell_shading(
                        cell,
                        "FDECEA" if values[0] == "FlowState-R1" else
                        (PALE_GREEN if improved else (WHITE if row_idx % 2 else PALE_GRAY)),
                    )
                    if improved or values[0] == "FlowState-R1":
                        run.bold = True
                    if col_idx == 5:
                        run.font.color.rgb = POSITIVE if improved else NEGATIVE
                else:
                    fills = {
                        "M4": PALE_GREEN,
                        "Electricity": PALE_BLUE,
                        "COVID Deaths": PALE_GRAY,
                        "Seattle Traffic": "FDECEA",
                    }
                    set_cell_shading(
                        cell, fills.get(values[0], WHITE if row_idx % 2 else PALE_GRAY)
                    )
                    if values[0] in fills:
                        run.bold = True
                    if col_idx == 1:
                        run.font.color.rgb = (POSITIVE if values[col_idx].startswith("-")
                                              else NEGATIVE)
        if row_idx == 0:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

    after = doc.add_paragraph()
    after.add_run("\u200b")
    after.paragraph_format.line_spacing = Pt(2)
    after.paragraph_format.space_after = Pt(4)


def process_text_between(doc: Document, chunk: str) -> None:
    for paragraph in re.split(r"\n\s*\n", chunk):
        add_body_paragraph(doc, paragraph)


def build() -> None:
    tex = strip_comments(SOURCE.read_text())
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("5  Zero-shot context selection with controllable risk")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph(style="Review Subtitle")
    subtitle.add_run("Revised review copy  ·  Expected-risk predictor and selectable harm profiles")

    token = re.compile(
        r"\\section\{[^}]+\}|"
        r"\\subsection\{[^}]+\}|"
        r"\\paragraph\{[^}]+\}|"
        r"\\begin\{equation\}(?:.*?)\\end\{equation\}|"
        r"\\begin\{figure\*?\}(?:.*?)\\end\{figure\*?\}|"
        r"\\begin\{table\*?\}(?:.*?)\\end\{table\*?\}",
        flags=re.DOTALL,
    )

    subsection_number = 0
    position = 0
    for match in token.finditer(tex):
        process_text_between(doc, tex[position:match.start()])
        block = match.group(0)
        if block.startswith("\\section"):
            pass  # represented by the manuscript title above
        elif block.startswith("\\subsection"):
            subsection_number += 1
            heading = plain(extract_braced(block, "subsection"))
            doc.add_heading(f"5.{subsection_number}  {heading}", level=2)
        elif block.startswith("\\paragraph"):
            heading = plain(extract_braced(block, "paragraph"))
            doc.add_heading(heading, level=3)
        elif block.startswith("\\begin{equation"):
            add_equation(doc, block)
        elif block.startswith("\\begin{figure"):
            add_figure(doc, block)
        elif block.startswith("\\begin{table"):
            add_table(doc, block)
        position = match.end()
    process_text_between(doc, tex[position:])

    # Remove empty paragraphs introduced only by source labels.
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip() and not paragraph._p.xpath(".//w:drawing"):
            paragraph._element.getparent().remove(paragraph._element)

    core = doc.core_properties
    core.title = "Section 5 — Zero-shot context selection with controllable risk"
    core.subject = "PredictCSL manuscript review copy"
    core.author = "PredictCSL"
    core.keywords = "time-series foundation models, context selection, risk calibration, abstention"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
